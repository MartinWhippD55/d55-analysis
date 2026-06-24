"""Generate Word documents from draft response markdown files."""
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def md_to_docx(md_path, docx_path, title):
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Skip the markdown title line and horizontal rules
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip markdown metadata
        if line.startswith('# Draft Response') or line == '---':
            i += 1
            continue
        
        # Empty line
        if not line:
            i += 1
            continue
        
        # Heading 2
        if line.startswith('## '):
            heading_text = line[3:]
            doc.add_heading(heading_text, level=2)
            i += 1
            continue
        
        # Heading 3
        if line.startswith('### '):
            heading_text = line[4:].strip('"')
            doc.add_heading(heading_text, level=3)
            i += 1
            continue
        
        # Table
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # Parse table
            header = [c.strip() for c in table_lines[0].split('|')[1:-1]]
            # Skip separator line
            data_rows = []
            for tl in table_lines[2:]:
                row = [c.strip() for c in tl.split('|')[1:-1]]
                data_rows.append(row)
            
            table = doc.add_table(rows=1 + len(data_rows), cols=len(header))
            table.style = 'Table Grid'
            
            for ci, cell_text in enumerate(header):
                cell_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
                table.rows[0].cells[ci].text = cell_text
            
            for ri, row in enumerate(data_rows):
                for ci, cell_text in enumerate(row):
                    cell_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
                    table.rows[ri + 1].cells[ci].text = cell_text
            
            continue
        
        # Numbered list
        if re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = text.replace('`', '')
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Bullet point
        if line.startswith('- '):
            text = line[2:]
            # Handle bold prefix like "**review-pr skill** - description"
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            text = text.replace('`', '')
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        
        # Bold paragraph header like **1. We are reviewing our code.**
        if line.startswith('**') and '**' in line[2:]:
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
            text = text.replace('`', '')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            i += 1
            continue
        
        # Regular paragraph
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
        text = text.replace('`', '')
        doc.add_paragraph(text)
        i += 1
    
    doc.save(docx_path)
    print(f"Created: {docx_path}")


if __name__ == '__main__':
    md_to_docx(
        'analysis/ESG/debt-management/draft-response-2.md',
        'analysis/ESG/debt-management/D55-Response-Detailed.docx',
        'D55 Response to Code Quality Concerns (Detailed)'
    )
    md_to_docx(
        'analysis/ESG/debt-management/draft-response-final.md',
        'analysis/ESG/debt-management/D55-Response-Executive.docx',
        'D55 Response to Code Quality Concerns (Executive)'
    )
